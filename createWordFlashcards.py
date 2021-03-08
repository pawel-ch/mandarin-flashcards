# -*- coding: utf-8 -*-
"""
    To run, have FlashcardTemplate.docx and flashcard-input.txt in the script directory. 
    flashcard-input.txt should have a list of Chinese characters/words with optional example sentences, 
    one entry per line, the characters separated from the example with white space (tab, space, etc) like this:
   
    高   大人在高高的山。
    本來
    悟空  悟空打了個妖怪。
    
    You can also include an English translation before the example, always in parentheses, like this:
    跟屁蟲 (someone's shadow)我的妹妹是跟屁蟲
    早晨  (early morning)早晨沒有人欣賞。
   
    The output is a .docx file with just the characters/words on one side and pinyin + example on the other side.
    The formatting is meant for printing on pre-perforating business card paper. 
"""
import codecs
import datetime
import re
from itertools import chain, islice

import pinyin
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

CHAR_COUNT_TO_STYLE = {
    1: "Hanzi Style 115",
    2: "Hanzi Style 100",
    3: "Hanzi Style 80",
    4: "Hanzi Style 60",
    5: "Hanzi Style 45",
}


def chunks(iterable, size=10):
    iterator = iter(iterable)
    for first in iterator:
        yield chain([first], islice(iterator, size - 1))


def process_page(doc, doc_table, terms, offset):
    print(f"Processing word list of length={len(terms)}")
    columns_num = 2

    for i, term in enumerate(terms, start=0):
        # hanzi
        (row, column) = (int((i + offset) / columns_num), i % columns_num)
        hanzi_p = doc_table.cell(row, column).paragraphs[0]
        hanzi_p.text = term["hanzi"]
        hanzi_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        char_count = len(term["hanzi"])
        new_style = CHAR_COUNT_TO_STYLE[char_count]
        hanzi_p.style = new_style

        # pinyin
        (row, column) = (int((i + offset) / columns_num) + 5, (i + 1) % columns_num)
        pinyin_cell = doc_table.cell(row, column)
        pinyin_p = pinyin_cell.add_paragraph(
            term["pinyin"] + " ", style=doc.styles["Pinyin"]
        )
        pinyin_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if term["english"]:
            pinyin_p.add_run(term["english"], style=doc.styles["English"])

        if term["example"]:
            para = pinyin_cell.add_paragraph(term["example"], style=doc.styles["Example"])
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def get_doc():
    doc = Document("FlashcardTemplate.docx")

    hanzi_style = doc.styles.add_style("Hanzi Style", WD_STYLE_TYPE.PARAGRAPH)
    hanzi_style.font.name = "DFKai-SB"
    hanzi_style_115 = doc.styles.add_style("Hanzi Style 115", WD_STYLE_TYPE.PARAGRAPH)
    hanzi_style_115.font.size = Pt(115)
    hanzi_style_115.base_style = doc.styles["Hanzi Style"]
    hanzi_style_100 = doc.styles.add_style("Hanzi Style 100", WD_STYLE_TYPE.PARAGRAPH)
    hanzi_style_100.font.size = Pt(100)
    hanzi_style_100.base_style = doc.styles["Hanzi Style"]
    hanzi_style_80 = doc.styles.add_style("Hanzi Style 80", WD_STYLE_TYPE.PARAGRAPH)
    hanzi_style_80.font.size = Pt(80)
    hanzi_style_80.base_style = doc.styles["Hanzi Style"]
    hanzi_style_60 = doc.styles.add_style("Hanzi Style 60", WD_STYLE_TYPE.PARAGRAPH)
    hanzi_style_60.font.size = Pt(60)
    hanzi_style_60.base_style = doc.styles["Hanzi Style"]
    hanzi_style_45 = doc.styles.add_style("Hanzi Style 45", WD_STYLE_TYPE.PARAGRAPH)
    hanzi_style_45.font.size = Pt(45)
    hanzi_style_45.base_style = doc.styles["Hanzi Style"]
    pinyin_style = doc.styles.add_style("Pinyin", WD_STYLE_TYPE.PARAGRAPH)
    pinyin_style.base_style = doc.styles["Normal"]
    pinyin_style.font.size = Pt(25)

    english_style = doc.styles.add_style("English", WD_STYLE_TYPE.CHARACTER)
    english_style.base_style = doc.styles["Normal"]
    english_style.font.size = Pt(15)

    example_style = doc.styles.add_style("Example", WD_STYLE_TYPE.PARAGRAPH)
    example_style.base_style = doc.styles["Hanzi Style"]
    example_style.font.size = Pt(20)

    return doc


def main():
    pinyin_list = []
    regex = re.compile(r"^(?P<hanzi>\S+)\s*(\((?P<english>.+)\))?\s*(?P<example>.+)?$")
    doc = get_doc()
    doc_table = doc.tables[0]

    with codecs.open("flashcard-input.txt", "r", "utf-8-sig") as fp:
        for line in fp.readlines():
            match = regex.match(line.strip())
            if match:
                term_info = match.groupdict()
                term_info["pinyin"] = pinyin.get(term_info["hanzi"])
                pinyin_list.append(term_info)

    rows_in_2_pages = 10

    for page_index, chunk in enumerate(chunks(pinyin_list, size=rows_in_2_pages), start=0):
        offset = page_index * rows_in_2_pages * 2

        # extend table for additional pages
        if offset:
            for i in range(rows_in_2_pages):
                new_row = doc_table.add_row()
                new_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                new_row.height = Inches(2)
                for new_cell in new_row.cells:
                    new_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        process_page(doc, doc_table, list(chunk), offset)

    time_str = str(datetime.datetime.now().strftime("%Y%m%d_%H-%M-%S"))
    output_doc_name = "flashcards-" + time_str + ".docx"
    doc.save(output_doc_name)


if __name__ == "__main__":
    main()
